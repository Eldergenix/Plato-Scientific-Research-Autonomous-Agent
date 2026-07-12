#!/usr/bin/env python3
"""Build the bioRxiv main manuscript and supplement as styled DOCX files.

Design system: ``narrative_proposal`` with a named ``academic_manuscript``
override (Times New Roman, black hierarchy, compact title block) and the
``editorial_cover`` first-page pattern adapted for a scientific preprint.
"""

from __future__ import annotations

import csv
import hashlib
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PREPRINT = ROOT / "preprint"
SKILL_SCRIPTS = Path(
    "/Users/0xnexis/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.709.11516/skills/documents/scripts"
)

sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLACK = RGBColor(0x18, 0x1B, 0x1F)
MUTED = RGBColor(0x55, 0x5F, 0x69)
BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"


def set_font(run, *, size=11, bold=None, italic=None, color=BLACK, name="Times New Roman"):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="B9C2CC", size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, before, after in (
        ("Heading 1", 15, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, size=9, color=MUTED, name="Arial")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def new_numbering_instance(doc: Document) -> int:
    """Create a fresh decimal numbering instance starting at one."""
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    source_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def configure_sections(doc: Document, running_title: str) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.42)
        section.footer_distance = Inches(0.42)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.text = running_title
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_font(p.runs[0], size=8.5, color=MUTED, name="Arial")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        add_page_field(footer_p)


def add_inline_markdown(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|https?://\S+)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, italic=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=9.5, name="Courier New")
        else:
            run = paragraph.add_run(token.rstrip(".,;"))
            set_font(run, color=BLUE)
            run.underline = True
            trailing = token[len(token.rstrip(".,;")) :]
            if trailing:
                tail = paragraph.add_run(trailing)
                set_font(tail)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_font(run)


def table_weights(rows: list[list[str]]) -> list[float]:
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        longest = max(len(row[index]) for row in rows)
        weights.append(min(max(longest, 6), 28))
    return weights


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    table.rows[0]._tr.trPr.append(tbl_header)

    for row_index, row in enumerate(rows):
        table.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
            set_cell_borders(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index < 3 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            set_font(run, size=8.3, bold=row_index == 0, name="Arial")

    widths = column_widths_from_weights(table_weights(rows), total_width_dxa=9792)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9792,
        indent_dxa=110,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 110, "end": 110},
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    rows = [raw[0]] + raw[2:] if len(raw) >= 2 else raw
    return rows, index


def add_title_block(doc: Document, title: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_font(run, size=18, bold=True, color=BLACK, name="Arial")


def build_from_markdown(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc, "Plato-Bio | bioRxiv preprint")

    index = 0
    first_heading = True
    active_numbering_id: int | None = None
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("# ") and first_heading:
            add_title_block(doc, stripped[2:].strip())
            first_heading = False
            index += 1
            continue
        if stripped.startswith("## "):
            active_numbering_id = None
            doc.add_heading(stripped[3:].strip(), level=1)
            index += 1
            continue
        if stripped.startswith("### "):
            active_numbering_id = None
            doc.add_heading(stripped[4:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("| "):
            active_numbering_id = None
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            active_numbering_id = None
            caption, relative = image_match.groups()
            image_path = source.parent / relative
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            shape = run.add_picture(str(image_path), width=Inches(6.25))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption.split(".", maxsplit=1)[0])
            caption_p = doc.add_paragraph(style="Caption")
            add_inline_markdown(caption_p, caption)
            index += 1
            continue

        if stripped == "<!-- PAGE BREAK -->":
            active_numbering_id = None
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            if active_numbering_id is None:
                active_numbering_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_numbering_id)
            spacer = p.add_run(" ")
            set_font(spacer)
            add_inline_markdown(p, number_match.group(1))
            index += 1
            continue
        if stripped.startswith("- "):
            active_numbering_id = None
            p = doc.add_paragraph(style="List Bullet")
            spacer = p.add_run(" ")
            set_font(spacer)
            add_inline_markdown(p, stripped[2:])
            index += 1
            continue

        active_numbering_id = None
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("- ")
                or re.match(r"^\d+\.\s+", candidate)
                or candidate.startswith("![")
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        text = " ".join(paragraph_lines).replace("  ", " ")
        p = doc.add_paragraph()
        if text.startswith("**Stefan Creadore"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
        elif text.startswith("¹ Eldergenix"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
        elif text.startswith("**Article category") or text.startswith("**bioRxiv") or text.startswith("**Keywords"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
        add_inline_markdown(p, text)

    core = doc.core_properties
    core.title = lines[0].removeprefix("# ")
    core.author = "Stefan Creadore"
    core.subject = "Verification-first computational biology research workflow"
    core.keywords = "bioinformatics, scientific agents, reproducibility, AlphaFold"
    doc.save(output)


def sha256(path: Path) -> str:
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return digest


def add_simple_paragraph(doc, text, *, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def build_supplement(output: Path) -> None:
    globin_dir = PREPRINT / "results" / "globin_benchmark"
    validation = json.loads((PREPRINT / "results" / "software_validation.json").read_text())
    targets = list(csv.DictReader((globin_dir / "target_summary.csv").open()))
    residues = list(csv.DictReader((globin_dir / "residue_metrics.csv").open()))

    doc = Document()
    configure_styles(doc)
    configure_sections(doc, "Plato-Bio | Supplementary Material")
    add_title_block(doc, "Supplementary Material: Plato-Bio")
    add_simple_paragraph(
        doc,
        "A reproducibility inventory for the structural-biology case study and deterministic software validation.",
        italic=True,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("S1. Reproduction commands", level=1)
    for command in (
        ".venv/bin/python preprint/experiments/run_globin_structure_benchmark.py",
        ".venv/bin/python preprint/experiments/run_software_validation.py",
        ".venv/bin/python preprint/experiments/build_summary_figures.py",
    ):
        p = doc.add_paragraph(style="List Bullet")
        spacer = p.add_run(" ")
        set_font(spacer)
        run = p.add_run(command)
        set_font(run, size=9, name="Courier New")

    doc.add_heading("S2. Structural input inventory", level=1)
    input_rows = [["File", "Bytes", "SHA-256"]]
    for path in sorted((globin_dir / "raw").glob("*.pdb")):
        input_rows.append([path.name, str(path.stat().st_size), sha256(path)])
    add_table(doc, input_rows)

    doc.add_heading("S3. Target-level structural results", level=1)
    target_rows = [[
        "Target", "Matched", "Identity", "RMSD Å", "Median Å", "≤2 Å", "ρ", "P"
    ]]
    for row in targets:
        target_rows.append([
            row["target"],
            row["matched_residues"],
            f"{float(row['sequence_identity']):.3f}",
            f"{float(row['ca_rmsd_angstrom']):.3f}",
            f"{float(row['median_ca_error_angstrom']):.3f}",
            f"{float(row['fraction_within_2a']):.3f}",
            f"{float(row['spearman_plddt_vs_negative_error']):.3f}",
            f"{float(row['spearman_pvalue']):.3g}",
        ])
    add_table(doc, target_rows)
    add_simple_paragraph(
        doc,
        f"The residue-level file contains {len(residues)} matched positions. It is supplied as residue_metrics.csv rather than reproduced as a {len(residues)}-row PDF table.",
        italic=True,
        size=9.5,
    )

    doc.add_heading("S4. Deterministic validation results", level=1)
    validation_rows = [["Suite", "Tests", "Passed", "Skipped", "Failures", "Errors", "Wall s"]]
    for name, suite in validation["suites"].items():
        passed = suite["tests"] - suite["skipped"] - suite["failures"] - suite["errors"]
        validation_rows.append([
            name,
            str(suite["tests"]),
            str(passed),
            str(suite["skipped"]),
            str(suite["failures"]),
            str(suite["errors"]),
            f"{suite['wall_seconds']:.2f}",
        ])
    add_table(doc, validation_rows)
    add_simple_paragraph(
        doc,
        "Targeted suites overlap with the full Python suite. Counts are not additive. Skips remain unvalidated live or platform-specific paths.",
        italic=True,
        size=9.5,
    )

    doc.add_heading("S5. Measurement-repair acceptance criteria", level=1)
    criteria = [
        "A GoldenTask with domain=biology constructs Plato with domain=biology.",
        "Method-signal recall is computed from methods.md and appears in summary metrics.",
        "Evidence JSONL persists every drafted Claim row before emitted EvidenceLink rows.",
        "A drafted claim with no supporting source remains in the denominator and yields unsupported_claim_rate=1.0.",
        "A synthetic rigid rotation and translation is recovered by the Kabsch implementation within numerical tolerance.",
    ]
    for item in criteria:
        p = doc.add_paragraph(style="List Bullet")
        spacer = p.add_run(" ")
        set_font(spacer)
        add_inline_markdown(p, item)

    doc.add_heading("S6. Interpretation and unavailable live lanes", level=1)
    add_simple_paragraph(
        doc,
        "The default evaluation runner executes idea and method stages only. It does not execute results or paper generation in this release. Live LLM, E2B, Modal, hosted PostgreSQL, and authenticated Hugging Face evaluations require external credentials or services and were not inferred from deterministic tests. Same-model reviewer roles are self-critique, not peer review. The autonomous-loop adapters do not establish autonomous scientific improvement.",
    )

    doc.add_heading("S7. bioRxiv packaging notes", level=1)
    notes = [
        "Main manuscript is supplied as one PDF with embedded figures and tables.",
        "Supplemental data are separate from the main manuscript.",
        "The work is an English-language computational methods/validation article with Results and Discussion.",
        "No clinical, human-subject, animal, or identifiable-person data are included.",
        "Generative AI assistance is disclosed; the human author remains accountable and the AI is not an author.",
        "The related Denario arXiv preprint is cited as prior work and is not duplicated by this study.",
    ]
    for item in notes:
        p = doc.add_paragraph(style="List Bullet")
        spacer = p.add_run(" ")
        set_font(spacer)
        add_inline_markdown(p, item)

    core = doc.core_properties
    core.title = "Supplementary Material: Plato-Bio"
    core.author = "Stefan Creadore"
    core.subject = "Reproducibility inventory"
    doc.save(output)


def main() -> int:
    main_docx = PREPRINT / "Plato-Bio-bioRxiv-preprint.docx"
    supplement_docx = PREPRINT / "Plato-Bio-Supplement.docx"
    build_from_markdown(PREPRINT / "manuscript.md", main_docx)
    build_supplement(supplement_docx)
    print(main_docx)
    print(supplement_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
